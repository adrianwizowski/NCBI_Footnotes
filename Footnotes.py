from contextlib import redirect_stdout
import sys
from Bio import Entrez
from docx import Document
import tkinter as tk
from tkinter.filedialog import askopenfilename, asksaveasfilename

class NCBI():
    def __init__(self):
        self.id_list = []
        self.parsed_urls = {}
        self.results = {}

    def parse_docx(self):
        '''Function reads Word file, searching for urls to pubmed database and creates a footnotes from that urls.
        No hyperlinks allowed.
        For example url:
        <footnote number> https://www.ncbi.nlm.nih.gov/pubmed/26087744
        will be changed to:
        <footnote number> Grace RF., Zanella A., Neufeld EJ., Morton DH., Eber S., Yaish H., Glader B. (2015) Erythrocyte pyruvate kinase
        deficiency: 2015 status report. Am. J. Hematol. 90(9): 825-30

        Function leaves leading footnotes number untouched.
        Saves output file with a .docx extension.'''

        tk_window = tk.Tk()
        tk_window.withdraw()
        chosen_file = askopenfilename(filetypes=[("Word Files", "*.docx"), ("Word Files", "*.doc")])
        if chosen_file == "":
            print("File Not Found")
            return None
        else:
            self.document = Document(chosen_file)
            #Searching links in file paragarphs
            for line in self.document.paragraphs:
                if line.text != '':
                    if "www.ncbi.nlm.nih.gov/pubmed" in line.text:
                        line_to_list = line.text.split('/')
                        for item in line_to_list:
                            if item.isdigit():
                                #Creating list of pubmed articles id's
                                id = item
                                self.id_list.append(id)
                                self.parsed_urls[line.text] = id

        self.create_footnotes()
        for key, value in self.parsed_urls.items():
            #Changing ulrs to a footnotes, leaving leading numbers untouched
            for line in self.document.paragraphs:
                if line.text == key:
                    footnote = self.parsed_urls[key][0]
                    if line.text[0:3].isdigit():
                        footnote = line.text[:3] + ' ' + footnote
                        line.text = footnote
                    elif line.text[0:2].isdigit():
                        footnote = line.text[:2] + ' ' + footnote
                        line.text = footnote
                    else:
                        footnote = line.text[:1] + ' ' + footnote
                        line.text = footnote

        self.document.save(asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")]))
        #Saves file as user choosen name



    def fetch_details(self,id_list):
        '''Sends list of pubmed articles ids to pubmed database using bio.python module.'''
        try:
            ids = ','.join(id_list)
            Entrez.email = 'your.email@example.com'
            handle = Entrez.efetch(db='pubmed',
                                   retmode='xml',
                                   id=ids)
            self.results = Entrez.read(handle)
        except RuntimeError:
            print("File doesn't contain any pubmed urls or urls are broken.")
            sys.exit(1)

    def create_footnotes(self):
        '''Creating footnotes from data received from pubmed database.
        Footnotes will be created using pattern:
        List of authors: AuthorLastName AuthorInitial., comma separated, (publication date in brackets), full title,
        ISO standardized journal title, The volume number of the journal in which the article was published,
        issue - part or supplement of the journal in which the article was published (if available) in brackets,
        International Standard Serial Number (if available)
        Example:
            Mukhopadhyay P., Mukherjee S., Ahsan K., Bagchi A., Pacher P., Das DK. (2010)
            Restoration of altered microRNA expression in the ischemic heart with resveratrol. PLoS ONE 5(12): e15705
        '''
        self.fetch_details(self.id_list)

        for i in self.results['PubmedArticle']:
            authors = []
            try:
                date = ' ('+ i['MedlineCitation']['Article']['Journal']['JournalIssue']['PubDate']['Year'] +') '
            except:
                date = ''
            title = i['MedlineCitation']['Article']['ArticleTitle']
            published = i['MedlineCitation']['Article']['Journal']['ISOAbbreviation']
            volume = i['MedlineCitation']['Article']['Journal']['JournalIssue']['Volume']
            try:
                issue = '('+ i['MedlineCitation']['Article']['Journal']['JournalIssue']['Issue'] +')'
            except:
                issue = ''

            try:
                issn = i['MedlineCitation']['Article']['Pagination']['MedlinePgn']
            except:
                issn = ''

            PMID = i['MedlineCitation']['PMID']
            for author in i['MedlineCitation']['Article']['AuthorList']:
                authors.append(author['LastName']+ ' ' +author['Initials']+'.')
            for k, v in self.parsed_urls.items():
                if PMID == v:
                    self.parsed_urls[k] = [', '.join(authors) + date + title +' '+ published +' '+ volume + issue + ': ' + issn, PMID]

    def get_ids(self):
        '''Creates list of pubmed article ids from parsed txt file.'''
        tk_window = tk.Tk()
        tk_window.withdraw()
        chosen_file = askopenfilename(filetypes=[("Text files", "*.txt")])
        if chosen_file == "":
            print("File 1 Not Found")
        else:
            with open(chosen_file) as input:
                for line in input:
                    if "www.ncbi.nlm.nih.gov/pubmed" in line:
                        line_to_list = line.strip('\n').split('/')
                        for item in line_to_list:
                            if item.isdigit():
                                # Creating list of pubmed articles id's
                                id = item
                                self.id_list.append(id)
                                self.parsed_urls[line] = id

    def save_to_txt(self):
        '''Saves result to a txt file.'''
        with open(asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")]), 'w') as f:
            with redirect_stdout(f):
                for key, value in self.parsed_urls.items():
                    print(key + '\n' + value[0] + '\n')

    def parse_txt(self):
        '''Run function to create footnotes from pubmed urls saves in txt file.'''
        self.get_ids()
        self.create_footnotes()
        self.save_to_txt()


if __name__ == "__main__":
    program = NCBI()
    program.parse_txt()